"""
create_full_resume.py
Creates a .docx résumé with the two centered header lines, followed
by EDUCATION, TECHNICAL SKILLS, and PROJECTS / EXPERIENCES – all with the
formatting shown in previous examples.

Important additions in this version
-----------------------------------
•  Bullet helper understands **bold** markup inside the text.
•  Cleaner loops: one bullet paragraph per list item (no char-by-char loops).

Requires:  pip install python-docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as ALIGN, WD_TAB_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING as LS





skills = [
    ("Frontend Development:",  " React Native, Flutter, TypeScript/JavaScript, HTML/CSS"),
    ("Backend Development:",   " Python, C, C++, Node.js, Data Structures & Algorithms, Distributed Systems"),
    ("Systems & OS:",          " Linux, Linux Kernel Development, Multithreading/Processes, Sockets (TCP/UDP), Bash/Shell"),
    ("Database & Cloud:",      " MongoDB, Firebase, Redis, AWS ECS, Linear Programming, Resource Optimization"),
    ("Development Tools:",     " Git, REST APIs, Stripe API, Google Maps API, Make/CMake"),
    ("Relevant Courses:",      " Systems Programming (Linux), Computer Architecture"),
]

projects = [
    (
        "your projects\time",
        [
            "description1",
            "2",
            
        ],
    ),
    
]



















# ────────────────────────────────────────────────────────────────
#  Helpers
# ────────────────────────────────────────────────────────────────
def apply_format(
    pf,
    *,
    alignment=ALIGN.LEFT,
    first_line_indent=None,
    left_indent=0,
    right_indent=11.25,
    space_before=0,
    space_after=0,
    line_spacing=1.0,
):
    pf.alignment = alignment
    pf.left_indent  = Pt(left_indent)
    pf.right_indent = Pt(right_indent)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    pf.line_spacing_rule = LS.MULTIPLE
    pf.line_spacing      = line_spacing
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)

def add_run(p, text, *, size=13, bold=False, italic=False,
            underline=False, font=None, color_hex=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    if font:
        r.font.name = font
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    return r

def add_bullet(p_text, *, size=14, color_hex="#0D0F1A"):
    """
    Create one bullet paragraph. Accepts **bold** inline markup.
    """
    p = doc.add_paragraph(style="List Bullet")
    apply_format(p.paragraph_format, left_indent=9, first_line_indent=-18)
    parts = re.split(r"(\*\*[^\*]+\*\*)", p_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            add_run(p, part[2:-2], size=size, bold=True, color_hex=color_hex)
        else:
            add_run(p, part,        size=size, bold=False, color_hex=color_hex)

# ────────────────────────────────────────────────────────────────
#  Document setup
# ────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]

section.page_width  = Inches(8.5)   # 612 pt
section.page_height = Inches(11)    # 792 pt
section.left_margin   = Pt(31)      # 0.43 in
section.right_margin  = Pt(30)      # 0.42 in
section.top_margin    = Pt(21)      # 0.29 in
section.bottom_margin = Pt(14)      # 0.19 in
section.header_distance = Pt(36)
section.footer_distance = Pt(36)

right_margin_in_inches = section.page_width.inches - section.right_margin.inches

# ────────────────────────────────────────────────────────────────
#  Header
# ────────────────────────────────────────────────────────────────
hdr1 = doc.add_paragraph()
apply_format(hdr1.paragraph_format, alignment=ALIGN.CENTER, first_line_indent=-9)
add_run(hdr1, "name", size=17, bold=True)

hdr2 = doc.add_paragraph()
apply_format(hdr2.paragraph_format, alignment=ALIGN.CENTER, first_line_indent=-9)
add_run(hdr2, "contacts", size=13)

# ────────────────────────────────────────────────────────────────
#  EDUCATION
# ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
apply_format(p.paragraph_format, first_line_indent=-9)
add_run(p, "EDUCATION", size=13, bold=True, underline=True,
        font="Times New Roman", color_hex="#000000")

p = doc.add_paragraph()
apply_format(p.paragraph_format)
add_run(p, " ", size=11, bold=True)
add_run(p, "school1", size=13, bold=True)
add_run(p, ", Silicon Valley, CA", size=13)

p = doc.add_paragraph()
apply_format(p.paragraph_format)
p.paragraph_format.tab_stops.add_tab_stop(
    Inches(right_margin_in_inches), alignment=WD_TAB_ALIGNMENT.RIGHT)
add_run(p, "Master of Computer Science in progress\tMarch 2027", size=13)

p = doc.add_paragraph()
apply_format(p.paragraph_format)
add_run(p, "University of California, Irvine", size=13, bold=True)
add_run(p, ", Irvine, CA", size=13)

p = doc.add_paragraph()
apply_format(p.paragraph_format)
p.paragraph_format.tab_stops.add_tab_stop(
    Inches(right_margin_in_inches), alignment=WD_TAB_ALIGNMENT.RIGHT)
add_run(p, "Bachelor of Science in Informatics\tMarch 2024", size=13)

# ────────────────────────────────────────────────────────────────
#  TECHNICAL SKILLS
# ────────────────────────────────────────────────────────────────


p = doc.add_paragraph()
apply_format(p.paragraph_format, first_line_indent=-9, space_before=4)
add_run(p, "TECHNICAL SKILLS", size=13, bold=True, underline=True)

for heading, tail in skills:
    p = doc.add_paragraph()
    apply_format(p.paragraph_format)
    add_run(p, heading, bold=True)
    add_run(p, tail)

# ────────────────────────────────────────────────────────────────
#  PROJECTS / EXPERIENCES
# ────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
apply_format(p.paragraph_format, first_line_indent=-9, space_before=4)
add_run(p, "PROJECTS / EXPERIENCES", size=13, bold=True, underline=True)

for title, bullets in projects:
    p = doc.add_paragraph()
    apply_format(p.paragraph_format)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(right_margin_in_inches), alignment=WD_TAB_ALIGNMENT.RIGHT)
    add_run(p, title, size=13, bold=True, color_hex="#0D0F1A")

    for b in bullets:
        add_bullet(b)

# ────────────────────────────────────────────────────────────────
#  SAVE
# ────────────────────────────────────────────────────────────────
doc.save("../resume_name.docx")
print("✓ Created ../resume_name.docx")
